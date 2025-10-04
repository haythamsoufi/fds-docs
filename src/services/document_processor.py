"""Document processing service with incremental updates."""

import os
import hashlib
import asyncio
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
import aiofiles
import logging
from datetime import datetime
import time
import io
import re
import unicodedata

import pypdf
from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

from src.core.models import DocumentModel, ChunkModel, DocumentStatus, DocumentCreate, ChunkCreate
from src.core.config import settings
from src.core.database import get_db_session
from src.core.monitoring import monitoring_service
from .text_splitter import AdaptiveTextSplitter
from .ocr_service import ocr_service

logger = logging.getLogger(__name__)


class DocumentChangeHandler(FileSystemEventHandler):
    """File system event handler for document changes."""
    
    def __init__(self, processor: 'DocumentProcessor'):
        self.processor = processor
        
    def on_modified(self, event):
        if not event.is_directory and self._is_supported_file(event.src_path):
            # Schedule safely from watchdog thread into the processor's loop
            if self.processor.loop and not self.processor.loop.is_closed():
                asyncio.run_coroutine_threadsafe(
                    self.processor.process_changed_file(event.src_path),
                    self.processor.loop
                )
            
    def on_created(self, event):
        if not event.is_directory and self._is_supported_file(event.src_path):
            if self.processor.loop and not self.processor.loop.is_closed():
                asyncio.run_coroutine_threadsafe(
                    self.processor.process_new_file(event.src_path),
                    self.processor.loop
                )
            
    def on_deleted(self, event):
        if not event.is_directory and self._is_supported_file(event.src_path):
            if self.processor.loop and not self.processor.loop.is_closed():
                asyncio.run_coroutine_threadsafe(
                    self.processor.remove_document(event.src_path),
                    self.processor.loop
                )
            
    def _is_supported_file(self, filepath: str) -> bool:
        """Check if file is supported."""
        return any(filepath.lower().endswith(ext) for ext in settings.supported_formats)


class DocumentProcessor:
    """Advanced document processor with incremental updates."""
    
    def __init__(self):
        self.text_splitter = AdaptiveTextSplitter()
        self.observer: Optional[Observer] = None
        self.loop: Optional[asyncio.AbstractEventLoop] = None
        self.processing_files: set = set()  # Track files currently being processed
        self.file_debounce_times: dict = {}  # Track last processing time for debouncing
        
    async def start_watching(self, watch_path: str = None):
        """Start watching for file changes."""
        if watch_path is None:
            watch_path = settings.documents_path
            
        try:
            # Capture the current running loop to schedule tasks thread-safely
            self.loop = asyncio.get_running_loop()

            self.observer = Observer()
            event_handler = DocumentChangeHandler(self)
            self.observer.schedule(event_handler, watch_path, recursive=True)
            self.observer.start()
            logger.info(f"Started watching directory: {watch_path}")
        except Exception as e:
            logger.error(f"Failed to start file watching: {e}")
            logger.info("Continuing without file watching - documents can still be processed manually")
            # Don't raise the exception, allow the system to continue without file watching
        
    async def stop_watching(self):
        """Stop watching for file changes."""
        if self.observer:
            try:
                self.observer.stop()
                self.observer.join()
                logger.info("Stopped watching for file changes")
            except Exception as e:
                logger.error(f"Error stopping file watcher: {e}")
                # Continue anyway since we're shutting down
            
    async def compute_file_hash(self, filepath: str) -> str:
        """Compute SHA-256 hash of file."""
        hash_sha256 = hashlib.sha256()
        async with aiofiles.open(filepath, 'rb') as f:
            async for chunk in self._read_chunks(f):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
        
    async def _read_chunks(self, file_obj, chunk_size: int = 8192):
        """Read file in chunks."""
        while True:
            chunk = await file_obj.read(chunk_size)
            if not chunk:
                break
            yield chunk
            
    async def extract_text_content(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text content and metadata from document."""
        file_ext = Path(filepath).suffix.lower()
        metadata = {
            "file_extension": file_ext,
            "extraction_timestamp": datetime.utcnow().isoformat()
        }
        
        try:
            if file_ext == '.pdf':
                return await self._extract_pdf_content(filepath, metadata)
            elif file_ext == '.docx':
                return await self._extract_docx_content(filepath, metadata)
            elif file_ext == '.txt':
                return await self._extract_txt_content(filepath, metadata)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error extracting content from {filepath}: {e}")
            raise
            
    async def _extract_pdf_content(self, filepath: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from PDF file with OCR fallback for scanned PDFs."""
        content_parts = []
        
        async with aiofiles.open(filepath, 'rb') as f:
            pdf_content = await f.read()
            
        try:
            # Use pypdf for extraction (wrap bytes in BytesIO)
            reader = pypdf.PdfReader(io.BytesIO(pdf_content))
        except Exception as e:
            # Handle PDF parsing errors (like graphics-related issues)
            error_msg = str(e)
            if "invalid float value" in error_msg or "gray non-stroke color" in error_msg:
                logger.warning(f"PDF parsing error (likely graphics-related) for {filepath}: {error_msg}")
                logger.info("Attempting OCR fallback for problematic PDF")
                # Fall back to OCR extraction
                try:
                    return await ocr_service.extract_text_with_ocr(filepath)
                except Exception as ocr_error:
                    logger.error(f"OCR fallback also failed for {filepath}: {ocr_error}")
                    return "", {"error": f"PDF parsing failed: {error_msg}. OCR fallback failed: {ocr_error}"}
            else:
                raise e
        
        # Extract metadata safely - handle different pypdf versions
        pdf_metadata = {}
        if reader.metadata:
            try:
                # Try new pypdf format first
                if hasattr(reader.metadata, 'get'):
                    pdf_metadata = dict(reader.metadata)
                # Fallback to old format
                elif hasattr(reader.metadata, '_data'):
                    pdf_metadata = reader.metadata._data
                # If neither works, try to access common attributes
                else:
                    pdf_metadata = {
                        'title': getattr(reader.metadata, 'title', ''),
                        'author': getattr(reader.metadata, 'author', ''),
                        'subject': getattr(reader.metadata, 'subject', ''),
                        'creator': getattr(reader.metadata, 'creator', ''),
                        'producer': getattr(reader.metadata, 'producer', ''),
                        'creation_date': str(getattr(reader.metadata, 'creation_date', '')),
                        'modification_date': str(getattr(reader.metadata, 'modification_date', ''))
                    }
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata from {filepath}: {e}")
                pdf_metadata = {}
        
        # Extract text using pypdf first
        extracted_text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    content_parts.append(f"[Page {page_num + 1}]\n{text}")
            except Exception as e:
                error_msg = str(e)
                if "invalid float value" in error_msg or "gray non-stroke color" in error_msg:
                    logger.warning(f"Graphics-related error on page {page_num + 1} of {filepath}: {error_msg}")
                    logger.info(f"Skipping problematic page {page_num + 1}, will continue with other pages")
                else:
                    logger.warning(f"Error extracting page {page_num + 1} from {filepath}: {e}")
                continue
        
        extracted_text = "\n\n".join(content_parts)
        
        # Check if this is a scanned PDF (very little text extracted) and OCR is enabled
        if len(extracted_text.strip()) < 100 and settings.ocr_enabled:
            logger.info(f"PDF appears to be scanned, attempting OCR extraction: {filepath}")
            ocr_text, ocr_metadata = await ocr_service.extract_text_with_ocr(filepath)
            if ocr_text.strip():
                extracted_text = ocr_text
                metadata.update({
                    "extraction_method": "ocr",
                    "ocr_metadata": ocr_metadata
                })
            else:
                metadata.update({
                    "extraction_method": "pypdf",
                    "ocr_attempted": True,
                    "ocr_metadata": ocr_metadata
                })
        else:
            # Try to extract tables even from text-based PDFs
            try:
                tables, table_metadata = await ocr_service.extract_tables_from_pdf(filepath)
                if tables:
                    table_text = []
                    for table in tables:
                        table_text.append(f"[Table {table['table_number']} on Page {table['page']}]\n{table['text_representation']}")
                    if table_text:
                        extracted_text += "\n\n" + "\n\n".join(table_text)
                        metadata["tables_extracted"] = len(tables)
                        metadata["table_metadata"] = table_metadata
            except Exception as e:
                logger.warning(f"Table extraction failed: {e}")
        
        metadata.update({
            "page_count": len(reader.pages),
            "pdf_metadata": pdf_metadata,
            "extraction_method": metadata.get("extraction_method", "pypdf")
        })
        
        return extracted_text, metadata
        
    async def _extract_docx_content(self, filepath: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from DOCX file."""
        doc = DocxDocument(filepath)
        
        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables
        tables_content = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                tables_content.append("\n".join(table_text))
                
        metadata.update({
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "has_tables": len(doc.tables) > 0
        })
        
        content_parts = paragraphs + tables_content
        return "\n\n".join(content_parts), metadata
        
    async def _extract_txt_content(self, filepath: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from text file."""
        async with aiofiles.open(filepath, 'r', encoding='utf-8') as f:
            content = await f.read()
            
        metadata.update({
            "character_count": len(content),
            "line_count": content.count('\n') + 1
        })
        
        return content, metadata
        
    async def process_document(self, filepath: str) -> Optional[str]:
        """Process a single document."""
        start_time = time.time()
        file_extension = Path(filepath).suffix.lower()
        
        # Check if file is already being processed
        if filepath in self.processing_files:
            logger.info(f"File already being processed, skipping: {filepath}")
            return None
            
        # Debounce: check if file was processed recently (within 5 seconds)
        current_time = time.time()
        if filepath in self.file_debounce_times:
            if current_time - self.file_debounce_times[filepath] < 5.0:
                logger.info(f"File processed recently, debouncing: {filepath}")
                return None
        
        # Mark file as being processed
        self.processing_files.add(filepath)
        self.file_debounce_times[filepath] = current_time

        async with monitoring_service.track_performance("document_processing", "service"):
            try:
                # Check if file exists and is accessible
                if not os.path.exists(filepath):
                    logger.error(f"File not found: {filepath}")
                    return None

                file_stats = os.stat(filepath)
                if file_stats.st_size == 0:
                    logger.warning(f"Empty file, skipping: {filepath}")
                    await monitoring_service.track_document_processing(
                        status="skipped",
                        format_type=file_extension,
                        duration=time.time() - start_time
                    )
                    return None

                if settings.max_file_size > 0 and file_stats.st_size > settings.max_file_size:
                    logger.warning(f"File too large: {filepath} ({file_stats.st_size} bytes, max: {settings.max_file_size} bytes)")
                    await monitoring_service.track_document_processing(
                        status="failed",
                        format_type=file_extension,
                        duration=time.time() - start_time
                    )
                    return None

                # Compute file hash
                file_hash = await self.compute_file_hash(filepath)

                # Use retry logic for database operations
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        async with get_db_session() as session:
                            # Check if document already exists and is unchanged
                            existing_doc = await self._get_document_by_path(session, filepath)
                            if existing_doc and existing_doc.file_hash == file_hash:
                                logger.info(f"Document unchanged, skipping: {filepath}")
                                await monitoring_service.track_document_processing(
                                    status="skipped",
                                    format_type=file_extension,
                                    duration=time.time() - start_time
                                )
                                return str(existing_doc.id)

                            # Extract content
                            content, metadata = await self.extract_text_content(filepath)

                            if not content.strip():
                                logger.warning(f"No content extracted from: {filepath}")
                                await monitoring_service.track_document_processing(
                                    status="failed",
                                    format_type=file_extension,
                                    duration=time.time() - start_time
                                )
                                return None

                            # Normalize and clean text prior to splitting
                            content = self._normalize_text(content)

                            # Create or update document record
                            if existing_doc:
                                document_id = await self._update_document(
                                    session, existing_doc, file_hash, content, metadata
                                )
                            else:
                                document_id = await self._create_document(
                                    session, filepath, file_hash, file_stats.st_size, content, metadata
                                )

                            # Track successful processing
                            processing_duration = time.time() - start_time
                            await monitoring_service.track_document_processing(
                                status="completed",
                                format_type=file_extension,
                                duration=processing_duration
                            )

                            logger.info(f"Successfully processed document: {filepath}")
                            return document_id

                    except Exception as db_error:
                        if "database is locked" in str(db_error) and attempt < max_retries - 1:
                            wait_time = (attempt + 1) * 0.5  # Exponential backoff
                            logger.warning(f"Database locked, retrying in {wait_time}s (attempt {attempt + 1}/{max_retries}): {filepath}")
                            await asyncio.sleep(wait_time)
                            continue
                        elif "UNIQUE constraint failed" in str(db_error):
                            logger.info(f"Document already exists, skipping: {filepath}")
                            await monitoring_service.track_document_processing(
                                status="skipped",
                                format_type=file_extension,
                                duration=time.time() - start_time
                            )
                            return None
                        else:
                            raise db_error

            except Exception as e:
                # Track failed processing
                processing_duration = time.time() - start_time
                await monitoring_service.track_document_processing(
                    status="failed",
                    format_type=file_extension,
                    duration=processing_duration
                )

                await monitoring_service.log_event(
                    "document_processing_failed",
                    {
                        "filepath": filepath,
                        "file_size": file_stats.st_size if 'file_stats' in locals() else 0,
                        "error": str(e)
                    },
                    level="error"
                )

                logger.error(f"Error processing document {filepath}: {e}")
                await self._mark_document_failed(filepath, str(e))
                return None
            finally:
                # Always remove from processing files
                self.processing_files.discard(filepath)
            
    async def _get_document_by_path(self, session: AsyncSession, filepath: str) -> Optional[DocumentModel]:
        """Get document by filepath."""
        result = await session.execute(
            select(DocumentModel).where(DocumentModel.filepath == filepath)
        )
        return result.scalar_one_or_none()
        
    async def _create_document(
        self, 
        session: AsyncSession, 
        filepath: str, 
        file_hash: str, 
        file_size: int,
        content: str, 
        metadata: Dict
    ) -> str:
        """Create new document record."""
        filename = os.path.basename(filepath)
        mime_type = self._get_mime_type(filepath)
        
        # Create document
        document = DocumentModel(
            filename=filename,
            filepath=filepath,
            file_hash=file_hash,
            file_size=file_size,
            mime_type=mime_type,
            status=DocumentStatus.PROCESSING,
            extra_metadata=metadata,
            content_preview=content[:500] + "..." if len(content) > 500 else content
        )
        
        session.add(document)
        await session.flush()  # Get the ID
        
        # Create chunks
        chunks = await self.text_splitter.split_text(content, metadata)
        # Deduplicate near-duplicate chunks via SimHash
        dedup_chunks = self._deduplicate_chunks(chunks)
        chunk_models = []
        
        for i, chunk in enumerate(dedup_chunks):
            chunk_hash = hashlib.sha256(chunk.content.encode()).hexdigest()
            # Attach simhash to chunk metadata for traceability
            try:
                sim_hash = self._simhash64(chunk.content)
                chunk.metadata = chunk.metadata or {}
                chunk.metadata.update({
                    "simhash": f"{sim_hash:016x}"
                })
            except Exception:
                # Best-effort; continue if simhash computation fails
                pass
            chunk_model = ChunkModel(
                document_id=document.id,
                chunk_index=i,
                content=chunk.content,
                content_hash=chunk_hash,
                extra_metadata=chunk.metadata
            )
            chunk_models.append(chunk_model)
            
        session.add_all(chunk_models)
        
        # Update document status
        document.status = DocumentStatus.COMPLETED
        document.processed_at = datetime.utcnow()
        document.chunk_count = len(dedup_chunks)
        
        await session.commit()
        return str(document.id)
        
    async def _update_document(
        self, 
        session: AsyncSession, 
        document: DocumentModel, 
        file_hash: str,
        content: str, 
        metadata: Dict
    ) -> str:
        """Update existing document with incremental chunk diff (add/update/remove)."""
        # Update document metadata first
        document.file_hash = file_hash
        document.status = DocumentStatus.PROCESSING
        document.extra_metadata = metadata
        document.content_preview = content[:500] + "..." if len(content) > 500 else content
        document.updated_at = datetime.utcnow()

        # Compute new chunks
        base_chunks = await self.text_splitter.split_text(content, metadata)
        new_chunks = self._deduplicate_chunks(base_chunks)
        new_hashes_by_index = {
            i: hashlib.sha256(chunk.content.encode()).hexdigest()
            for i, chunk in enumerate(new_chunks)
        }

        # Load existing chunks
        existing_result = await session.execute(
            select(ChunkModel).where(ChunkModel.document_id == document.id)
        )
        existing_chunks = existing_result.scalars().all()
        existing_by_index = {c.chunk_index: c for c in existing_chunks}

        new_indexes = set(new_hashes_by_index.keys())
        existing_indexes = set(existing_by_index.keys())

        # Determine operations
        to_remove = sorted(existing_indexes - new_indexes)
        to_consider_update = sorted(existing_indexes & new_indexes)
        to_add = sorted(new_indexes - existing_indexes)

        # Remove chunks that no longer exist
        if to_remove:
            await session.execute(
                delete(ChunkModel).where(
                    ChunkModel.document_id == document.id,
                    ChunkModel.chunk_index.in_(to_remove)
                )
            )

        # Update modified chunks (by index, only if content hash changed)
        for idx in to_consider_update:
            existing_chunk = existing_by_index[idx]
            new_chunk = new_chunks[idx]
            new_hash = new_hashes_by_index[idx]
            if existing_chunk.content_hash != new_hash:
                existing_chunk.content = new_chunk.content
                existing_chunk.content_hash = new_hash
                existing_chunk.extra_metadata = new_chunk.metadata

        # Add new chunks
        for idx in to_add:
            new_chunk = new_chunks[idx]
            new_hash = new_hashes_by_index[idx]
            session.add(
                ChunkModel(
                    document_id=document.id,
                    chunk_index=idx,
                    content=new_chunk.content,
                    content_hash=new_hash,
                    extra_metadata=new_chunk.metadata
                )
            )

        # Finalize document status
        document.status = DocumentStatus.COMPLETED
        document.processed_at = datetime.utcnow()
        document.chunk_count = len(new_chunks)

        await session.commit()
        return str(document.id)
        
    async def _mark_document_failed(self, filepath: str, error_message: str):
        """Mark document as failed."""
        async with get_db_session() as session:
            document = await self._get_document_by_path(session, filepath)
            if document:
                document.status = DocumentStatus.FAILED
                document.extra_metadata = document.extra_metadata or {}
                document.extra_metadata["error"] = error_message
                await session.commit()
                
    def _get_mime_type(self, filepath: str) -> str:
        """Get MIME type from file extension."""
        ext = Path(filepath).suffix.lower()
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain'
        }
        return mime_types.get(ext, 'application/octet-stream')
        
    async def process_directory(self, directory_path: str = None) -> Dict[str, Any]:
        """Process all documents in a directory."""
        if directory_path is None:
            directory_path = settings.documents_path
            
        results = {
            "processed": 0,
            "failed": 0,
            "skipped": 0,
            "total": 0,
            "errors": []
        }
        
        try:
            # Check if directory exists
            if not os.path.exists(directory_path):
                logger.warning(f"Documents directory does not exist: {directory_path}")
                results["errors"].append(f"Directory not found: {directory_path}")
                return results
                
            # Find all supported files
            supported_files = []
            for root, _, files in os.walk(directory_path):
                for file in files:
                    if any(file.lower().endswith(ext) for ext in settings.supported_formats):
                        supported_files.append(os.path.join(root, file))

            # Cleanup database entries whose files are missing on disk
            try:
                async with get_db_session() as session:
                    db_paths_result = await session.execute(
                        select(DocumentModel.filepath)
                    )
                    db_filepaths = {row[0] for row in db_paths_result}
                    fs_filepaths = set(supported_files)
                    missing_paths = db_filepaths - fs_filepaths

                    if missing_paths:
                        # Find document ids for missing paths
                        missing_docs_result = await session.execute(
                            select(DocumentModel).where(DocumentModel.filepath.in_(list(missing_paths)))
                        )
                        missing_docs = missing_docs_result.scalars().all()
                        missing_ids = [doc.id for doc in missing_docs]

                        if missing_ids:
                            # Delete chunks for missing docs
                            await session.execute(
                                delete(ChunkModel).where(ChunkModel.document_id.in_(missing_ids))
                            )
                            # Delete documents
                            await session.execute(
                                delete(DocumentModel).where(DocumentModel.id.in_(missing_ids))
                            )
                            await session.commit()
                            logger.info(f"Cleaned up {len(missing_ids)} documents missing from filesystem")
            except Exception as cleanup_err:
                logger.error(f"Cleanup during directory processing failed: {cleanup_err}")
                        
            results["total"] = len(supported_files)
            logger.info(f"Found {len(supported_files)} supported files in {directory_path}")
            
            if not supported_files:
                logger.info("No supported files found to process")
                return results
            
            # Process files in batches
            batch_size = settings.batch_size
            for i in range(0, len(supported_files), batch_size):
                batch = supported_files[i:i + batch_size]
                logger.info(f"Processing batch {i//batch_size + 1}: {len(batch)} files")
                
                tasks = [self.process_document(filepath) for filepath in batch]
                batch_results = await asyncio.gather(*tasks, return_exceptions=True)
                
                for j, result in enumerate(batch_results):
                    filepath = batch[j]
                    if isinstance(result, Exception):
                        results["failed"] += 1
                        error_msg = f"Failed to process {filepath}: {str(result)}"
                        results["errors"].append(error_msg)
                        logger.error(error_msg)
                    elif result is None:
                        results["skipped"] += 1
                        logger.info(f"Skipped {filepath} (no content or already processed)")
                    else:
                        results["processed"] += 1
                        logger.info(f"Successfully processed {filepath}")
                        
        except Exception as e:
            error_msg = f"Error during directory processing: {str(e)}"
            results["errors"].append(error_msg)
            logger.error(error_msg)
                    
        logger.info(f"Directory processing complete: {results}")
        return results
        
    async def process_new_file(self, filepath: str):
        """Process newly created file."""
        logger.info(f"Processing new file: {filepath}")
        await self.process_document(filepath)
        
    async def process_changed_file(self, filepath: str):
        """Process changed file."""
        logger.info(f"Processing changed file: {filepath}")
        await self.process_document(filepath)
        
    async def remove_document(self, filepath: str):
        """Remove document from database."""
        logger.info(f"Removing document: {filepath}")
        async with get_db_session() as session:
            document = await self._get_document_by_path(session, filepath)
            if document:
                # Delete chunks first
                await session.execute(
                    delete(ChunkModel).where(ChunkModel.document_id == document.id)
                )
                # Delete document
                await session.delete(document)
                await session.commit()
                logger.info(f"Document removed from database: {filepath}")

    # ---- Deduplication utilities ----
    def _tokenize_for_simhash(self, text: str) -> List[str]:
        """Simple tokenization to lowercase words; can be improved later."""
        return re.findall(r"[A-Za-z0-9]+", text.lower())

    def _hash64(self, token: str) -> int:
        """Stable 64-bit hash using SHA-1 truncated to 64 bits."""
        import hashlib as _hashlib  # local alias to avoid confusion
        digest = _hashlib.sha1(token.encode("utf-8")).digest()
        # Take first 8 bytes as big-endian 64-bit int
        return int.from_bytes(digest[:8], byteorder="big", signed=False)

    def _simhash64(self, text: str) -> int:
        """Compute 64-bit SimHash for a piece of text."""
        bits = [0] * 64
        tokens = self._tokenize_for_simhash(text)
        if not tokens:
            return 0
        # Weight by term frequency
        from collections import Counter
        freq = Counter(tokens)
        for token, weight in freq.items():
            h = self._hash64(token)
            for i in range(64):
                if (h >> i) & 1:
                    bits[i] += weight
                else:
                    bits[i] -= weight
        # Construct hash: bit is 1 if positive
        value = 0
        for i in range(64):
            if bits[i] > 0:
                value |= (1 << i)
        return value

    def _hamming_distance(self, a: int, b: int) -> int:
        """Compute Hamming distance between two 64-bit integers."""
        x = a ^ b
        # Kernighan's algorithm
        count = 0
        while x:
            x &= x - 1
            count += 1
        return count

    def _deduplicate_chunks(
        self,
        chunks: List[Any],
        distance_threshold: int = 3,
        min_length: int = 40
    ) -> List[Any]:
        """Remove near-duplicate chunks using SimHash with Hamming distance threshold.

        - Skips chunks shorter than min_length from being considered duplicates (keeps them all)
        - Maintains original order
        - Annotates kept chunks' metadata with simhash
        """
        kept: List[Any] = []
        fingerprints: List[int] = []
        for chunk in chunks:
            text = (chunk.content or "").strip()
            if len(text) < min_length:
                kept.append(chunk)
                continue
            fp = self._simhash64(text)
            is_duplicate = any(self._hamming_distance(fp, prev) <= distance_threshold for prev in fingerprints)
            if is_duplicate:
                continue
            fingerprints.append(fp)
            # Attach simhash in metadata for kept chunks; others are dropped entirely
            try:
                chunk.metadata = chunk.metadata or {}
                chunk.metadata.update({"simhash": f"{fp:016x}"})
            except Exception:
                pass
            kept.append(chunk)
        return kept

    def _normalize_text(self, text: str) -> str:
        """Normalize and clean extracted text; remove boilerplate and artifacts.

        Operations:
        - Unicode normalize to NFKC
        - Normalize line endings to \n
        - Fix hyphenation across line breaks (e.g., "trans-\nform" -> "transform")
        - Remove control characters except tab/newline
        - Collapse excessive whitespace and newlines
        - Remove common boilerplate lines like "Page X of Y" (preserve "[Page N]" markers)
        """
        if not text:
            return text

        # Unicode normalization
        normalized = unicodedata.normalize("NFKC", text)

        # Standardize newlines
        normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")

        # Fix hyphenation across line breaks: word-\nword -> wordword (keep a single newline)
        normalized = re.sub(r"(\w)-\n(\w)", r"\1\2\n", normalized)

        # Remove control characters except tab/newline
        normalized = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", normalized)

        # Remove boilerplate lines like "Page X of Y" (not [Page N])
        lines = []
        for line in normalized.split("\n"):
            if re.match(r"^\s*Page\s+\d+(\s+of\s+\d+)?\s*$", line):
                continue
            lines.append(line)
        normalized = "\n".join(lines)

        # Collapse 3+ newlines to at most 2
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)

        # Collapse internal whitespace sequences to single spaces within lines
        normalized = "\n".join(re.sub(r"[ \t]{2,}", " ", ln).strip() for ln in normalized.split("\n"))

        return normalized
